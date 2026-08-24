"""
parser/resume_parser.py
Extracts raw text, fonts (PDF only), and contact info from PDF/DOCX resumes.
"""

import re
import pdfplumber
import docx

from parser.section_splitter import split_into_sections

STANDARD_FONTS = [
    # Traditional ATS fonts
    "arial", "calibri", "times new roman", "helvetica", "georgia", "cambria", "verdana",
    # LaTeX ATS-friendly fonts & internal font identifiers
    "latin modern", "lmroman", "lmsans", "lmmono",
    "tex gyre heros", "texgyreheros", "gyreheros",
    "tex gyre termes", "texgyretermes", "gyretermes",
    "libertinus", "libertine", "linlibertine"
]


def extract_text_and_fonts_from_pdf(file_path):
    """Extract text and font names from a PDF using pdfplumber."""
    text = ""
    fonts_found = set()

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
            for char in page.chars:
                font_name = char.get("fontname", "")
                if font_name:
                    # strip subset prefix like "ABCDEF+Arial"
                    fonts_found.add(font_name.split("+")[-1])

    return text, list(fonts_found)


def extract_text_from_docx(file_path):
    """Extract text and font names from a DOCX using python-docx."""
    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)

    fonts_found = set()
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font and run.font.name:
                fonts_found.add(run.font.name)

    return text, list(fonts_found)


def extract_text(file_path):
    """Detect file type by extension and extract text + fonts accordingly."""
    if file_path.lower().endswith(".pdf"):
        return extract_text_and_fonts_from_pdf(file_path)
    elif file_path.lower().endswith(".docx"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file type. Please upload PDF or DOCX.")


def get_contact_info(text):
    """
    Extract email, phone, and professional profile mentions.
    Note: PDF hyperlinks often only expose their visible label (e.g. "LinkedIn"),
    not the underlying URL, in extracted text — so we check for platform
    name mentions rather than requiring a full URL match.
    """
    email_match = re.search(r'[\w.-]+@[\w.-]+\.\w+', text)
    phone_match = re.search(
        r'(\+?\d{1,3}[-.\s]?)?\(?\d{3,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}', text)

    text_lower = text.lower()
    has_linkedin = bool(re.search(r'linkedin(\.com)?', text_lower))
    has_github = bool(re.search(r'github(\.com)?', text_lower))
    has_portfolio = bool(re.search(r'portfolio', text_lower))

    return {
        "email": email_match.group(0) if email_match else None,
        "phone": phone_match.group(0) if phone_match else None,
        "linkedin": has_linkedin,
        "github": has_github,
        "portfolio": has_portfolio,
    }


def check_font_compliance(fonts_found):
    """
    Checks whether detected fonts are ATS-standard.
    Returns (is_compliant, list_of_nonstandard_fonts_found).
    If no fonts detected (e.g. DOCX with no explicit run fonts, meaning
    it's using the document's default theme font), we treat it as compliant
    since most default Word themes use standard fonts.
    """
    if not fonts_found:
        return True, []

    nonstandard = [f for f in fonts_found if not any(
        std in f.lower() for std in STANDARD_FONTS)]
    is_compliant = len(nonstandard) == 0
    return is_compliant, nonstandard


def parse_resume(file_path):
    """
    Main entry point: takes a file path, returns a dict with
    raw text, sections, contact info, fonts, and word count.
    """
    raw_text, fonts_found = extract_text(file_path)
    sections = split_into_sections(raw_text)
    contact_info = get_contact_info(raw_text)
    font_compliant, nonstandard_fonts = check_font_compliance(fonts_found)

    return {
        "raw_text": raw_text,
        "sections": sections,
        "contact_info": contact_info,
        "word_count": len(raw_text.split()),
        "fonts_found": fonts_found,
        "font_compliant": font_compliant,
        "nonstandard_fonts": nonstandard_fonts,
    }
